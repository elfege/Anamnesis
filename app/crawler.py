"""
Anamnesis Crawler — Active source scanner and episode ingester.

Periodically scans known knowledge sources and ingests new content as episodes.
Sources are mounted read-only into the Docker container (or staged via rsync).
Deduplication is by SHA-256 hash of section content (crawl_state collection).

Source types:
  - Named SOURCES: specific files (genesis, handoff, history, intercom, user profile)
  - Project scanner: discovers projects by docker-compose.yml presence, ingests
    CLAUDE.md, README.md, docker-compose.yml, *.sh, *.py at project root
  - Deep project scanner: recursively scans all 0_* dirs + HUBITAT/NETWORK for
    code files (.ino, .cpp, .h, .groovy, .py, .sh, .js, .ts, etc.) up to 64KB each
  - Scripts scanner: crawls 0_SCRIPTS/ for .sh files (bash style authority)
  - Teachings: each .md file in 0_TEACHINGS/ = one episode
  - Documents: .docx files from OneDrive

Machines are configured via the dashboard (stored in MongoDB).
Remote machines are staged via sync_sources.sh (runs hourly via cron).
"""

import asyncio
import hashlib
import logging
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import zipfile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from lxml import etree

from database import get_episodes_collection, get_settings_collection
from embedding import get_embedding
from config import MONGO_DB

logger = logging.getLogger("anamnesis.crawler")

# ─── Source Config ────────────────────────────────────────────────
# All config lives in MongoDB (settings collection, _id: "crawler_config").
# On first run with empty DB, empty lists are seeded — configure via dashboard.
# No hardcoded paths in code (this repo is public).

# Runtime config — populated by load_crawler_config() from DB at startup.
SOURCES: list[dict] = []
MACHINE_ROOTS: dict[str, str] = {}


# ─── Persistent Config ───────────────────────────────────────────

async def load_crawler_config():
    """Load crawler config from MongoDB. Seeds empty config on first run.
    After this call, SOURCES and MACHINE_ROOTS are always DB-driven."""
    global SOURCES, MACHINE_ROOTS
    coll = get_settings_collection()
    doc = await coll.find_one({"_id": "crawler_config"})

    if not doc:
        # First run — seed empty config; user configures via dashboard
        logger.warning("No crawler config in DB — seeding empty. Configure via dashboard.")
        seed = {"sources": [], "machine_roots": {}}
        await coll.update_one({"_id": "crawler_config"}, {"$set": seed}, upsert=True)
        doc = seed

    SOURCES = doc.get("sources", [])
    MACHINE_ROOTS = doc.get("machine_roots", {})
    logger.info(f"Crawler config loaded: {len(SOURCES)} sources, {len(MACHINE_ROOTS)} machine roots")

    # Restore persisted stats (survive container restarts)
    stats_doc = await coll.find_one({"_id": "crawler_stats"})
    if stats_doc:
        _crawler_state["total_episodes_ingested"] = stats_doc.get("total_episodes_ingested", 0)
        _crawler_state["episodes_ingested_last_run"] = stats_doc.get("episodes_ingested_last_run", 0)
        _crawler_state["last_run"] = stats_doc.get("last_run")
        _crawler_state["last_run_duration_seconds"] = stats_doc.get("last_run_duration_seconds", 0)
        logger.info(f"Crawler stats restored: {_crawler_state['total_episodes_ingested']} total ingested")


async def save_crawler_config(sources: list = None, machine_roots: dict = None):
    """Persist crawler config to MongoDB. Only updates provided fields."""
    global SOURCES, MACHINE_ROOTS
    coll = get_settings_collection()
    update = {}
    if sources is not None:
        update["sources"] = sources
        SOURCES = sources
    if machine_roots is not None:
        update["machine_roots"] = machine_roots
        MACHINE_ROOTS = machine_roots
    if update:
        await coll.update_one({"_id": "crawler_config"}, {"$set": update}, upsert=True)


async def get_crawler_config() -> dict:
    """Return current crawler config for the API."""
    return {
        "sources": SOURCES,
        "machine_roots": MACHINE_ROOTS,
    }

# File extensions treated as code (ingested as whole-file episodes, not section-split)
_CODE_EXTENSIONS = {".sh", ".py", ".yml", ".yaml"}

# Extended code extensions for deep project scanning
_DEEP_CODE_EXTENSIONS = {
    ".ino", ".cpp", ".c", ".h",          # Arduino / C/C++
    ".groovy",                             # Hubitat
    ".py", ".sh",                          # Python / Bash
    ".js", ".ts", ".jsx", ".tsx",          # JS/TS
    ".yml", ".yaml",                       # Config
    ".java", ".go", ".rs",                 # Other langs
}

# Dirs to skip when recursively scanning projects
_SKIP_DIRS = {
    "build", "node_modules", ".git", "__pycache__", ".gradle",
    "libraries", "ARCHIVE", "DEPRECATED", "TRASH", "CONFLICTS",
    "venv", ".venv", "dist", "target", ".pio", ".platformio",
    "Copy", "copy", "BACKUP", ".idea", ".vscode",
    "Marlin", "marlin", "ArduinoJson", "ESP8266WiFi",        # Arduino libraries/firmware
    "packages", "vendor", "third_party", "deps", "lib",      # vendored deps
    "0_ANAMNESIS_SOURCES",                                     # staged copies (avoid double-ingestion)
}

# Project dirs to scan (beyond docker-compose-based discovery)
# Matches: any dir starting with "0_" plus these named dirs
_EXTRA_PROJECT_NAMES = {"HUBITAT", "0_HUBITAT", "NETWORK", "0_NETWORK"}

# Max file size for code ingestion (skip huge generated/vendored files)
_MAX_CODE_FILE_BYTES = 64 * 1024  # 64 KB

# Dir name fragments to skip when walking 0_SCRIPTS/
_SCRIPTS_SKIP = {"DEPRECATED", "TRASH", "ARCHIVE", "CONFLICTS"}

# Teachings directory — each .md file is a separate source
TEACHINGS_DIR = "/sources/teachings"

# Documents directory — .docx files from OneDrive
DOCUMENTS_DIR = "/sources/documents"

# ─── Crawler State ───────────────────────────────────────────────

_crawler_state = {
    "running": False,
    "last_run": None,
    "last_run_duration_seconds": 0,
    "episodes_ingested_last_run": 0,
    "total_episodes_ingested": 0,
    "errors": [],
    "current_activity": "",                                        # live one-liner for UI
}

_crawler_task: Optional[asyncio.Task] = None


def get_crawler_status() -> dict:
    """Return current crawler state for the dashboard/API."""
    return {**_crawler_state}


def _set_activity(msg: str):
    """Update the live activity one-liner shown in the UI."""
    _crawler_state["current_activity"] = msg


# ─── Parsing ─────────────────────────────────────────────────────

def _content_hash(text: str) -> str:
    """SHA-256 hash of text content for deduplication."""
    return hashlib.sha256(text.strip().encode("utf-8")).hexdigest()


def _parse_markdown_sections(text: str) -> list[dict]:
    """Split markdown text into sections by ### headers.

    Returns list of dicts with 'title', 'content', 'hash'.
    Skips sections that are too short to be meaningful (< 50 chars).
    """
    sections = []
    # Split on ### headers (keep the header as part of the section)
    parts = re.split(r"(?=^### )", text, flags=re.MULTILINE)

    for part in parts:
        part = part.strip()
        if not part or len(part) < 50:                             # skip tiny fragments
            continue

        # Extract title from ### header if present
        title_match = re.match(r"^###\s+(.+?)$", part, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else "untitled"

        sections.append({
            "title": title,
            "content": part,
            "hash": _content_hash(part),
        })

    return sections


def _parse_intercom_messages(text: str) -> list[dict]:
    """Parse intercom.md into individual messages by ### MSG-* headers."""
    sections = []
    parts = re.split(r"(?=^### MSG-)", text, flags=re.MULTILINE)

    for part in parts:
        part = part.strip()
        if not part or not part.startswith("### MSG-"):
            continue
        if len(part) < 30:                                         # skip ACK-only fragments
            continue

        # Extract MSG ID
        msg_match = re.match(r"^### (MSG-\d+)", part)
        msg_id = msg_match.group(1) if msg_match else "unknown"

        sections.append({
            "title": msg_id,
            "content": part,
            "hash": _content_hash(part),
        })

    return sections


def _parse_docx(filepath: str) -> str:
    """Extract all text from a .docx file. Returns full text as a string."""
    try:
        doc = DocxDocument(filepath)
    except PackageNotFoundError:
        logger.warning(f"Corrupt or unreadable docx: {filepath}")
        return ""
    except Exception as e:
        logger.warning(f"Failed to parse docx {filepath}: {e}")
        return ""

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _parse_pdf(filepath: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed — skipping PDF: %s", filepath)
        return ""
    try:
        pages = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        logger.warning(f"Failed to parse PDF {filepath}: {e}")
        return ""


def _parse_odt(filepath: str) -> str:
    """Extract text from an OpenDocument (.odt) file using odfpy."""
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P as OdfP
        from odf import teletype
    except ImportError:
        logger.warning("odfpy not installed — skipping ODT: %s", filepath)
        return ""
    try:
        doc = odf_load(filepath)
        paragraphs = []
        for p in doc.getElementsByType(OdfP):
            text = teletype.extractText(p).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"Failed to parse ODT {filepath}: {e}")
        return ""


def _parse_pages(filepath: str) -> str:
    """Extract text from Apple Pages (.pages) files.

    Pages files are ZIP archives. Text is in Index/Document.iwa (protobuf) which
    is hard to parse, but some Pages files also contain a preview PDF or an
    Index/Document.xml fallback.  We try multiple strategies:
      1. preview.pdf inside the archive (parsed via pdfplumber)
      2. index.xml (older iWork format — XML with paragraphs)
      3. buildVersionHistory.plist (has plain text fragments)
    """
    if not zipfile.is_zipfile(filepath):
        logger.debug(f"Not a valid Pages archive: {filepath}")
        return ""

    try:
        with zipfile.ZipFile(filepath) as zf:
            names = zf.namelist()

            # Strategy 1: embedded preview PDF
            pdf_names = [n for n in names if n.lower().endswith(".pdf")]
            if pdf_names:
                import tempfile
                with zf.open(pdf_names[0]) as pdf_entry:
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
                        tmp.write(pdf_entry.read())
                        tmp.flush()
                        text = _parse_pdf(tmp.name)
                        if text and len(text) > 30:
                            return text

            # Strategy 2: index.xml (older iWork XML format)
            xml_names = [n for n in names if n.lower().endswith(".xml")]
            for xml_name in xml_names:
                try:
                    with zf.open(xml_name) as xf:
                        tree = etree.parse(xf)
                        texts = tree.xpath("//text()")
                        joined = " ".join(t.strip() for t in texts if t.strip())
                        if len(joined) > 30:
                            return joined
                except Exception:
                    continue

            # Strategy 3: brute-force — read all text-like content
            text_fragments = []
            for name in names:
                if any(name.lower().endswith(ext) for ext in (".txt", ".strings")):
                    try:
                        text_fragments.append(zf.read(name).decode("utf-8", errors="replace"))
                    except Exception:
                        pass
            if text_fragments:
                combined = "\n".join(text_fragments)
                if len(combined) > 30:
                    return combined

    except Exception as e:
        logger.warning(f"Failed to parse Pages file {filepath}: {e}")
    return ""


def _parse_plain_text(filepath: str) -> str:
    """Read a plain text / markdown file."""
    try:
        return Path(filepath).read_text(encoding="utf-8", errors="replace").strip()
    except Exception as e:
        logger.warning(f"Failed to read text file {filepath}: {e}")
        return ""


# Supported document extensions → parser function
_DOC_PARSERS = {
    ".docx":  _parse_docx,
    ".pdf":   _parse_pdf,
    ".odt":   _parse_odt,
    ".pages": _parse_pages,
    ".md":    _parse_plain_text,
    ".txt":   _parse_plain_text,
    ".rtf":   _parse_plain_text,   # plain-text extraction (formatting lost)
}


async def _categorize_document(filename: str, content: str) -> list[str]:
    """Infer tags for a document file using patterns stored in MongoDB."""
    from database import load_doc_tag_patterns
    ext = os.path.splitext(filename)[1].lower()
    format_tag = ext.lstrip(".")  # e.g. "pdf", "docx", "pages"
    tags = ["document", "onedrive", format_tag]
    name_lower = filename.lower()
    patterns = await load_doc_tag_patterns()
    for p in patterns:
        subject = name_lower if p.get("field", "filename") == "filename" else content.lower()
        try:
            if p.get("regex"):
                flags = re.IGNORECASE if p.get("ignorecase") else 0
                matched = re.search(p["match"], subject, flags)
            else:
                matched = p["match"].lower() in subject if p.get("ignorecase") else p["match"] in subject
        except re.error:
            continue
        if matched and p["tag"] not in tags:
            tags.append(p["tag"])
    return tags


def _parse_code_file(filepath: str) -> list[dict]:
    """Read a code/config file as a single section (no header splitting)."""
    try:
        content = Path(filepath).read_text(encoding="utf-8", errors="replace").strip()
    except Exception:
        return []
    if len(content) < 20:
        return []
    return [{"title": Path(filepath).name, "content": content, "hash": _content_hash(content)}]


def _make_episode_id(source_name: str, section_title: str, content_hash: str) -> str:
    """Generate a deterministic episode_id from source + content."""
    # Clean title for use in ID
    clean_title = re.sub(r"[^a-zA-Z0-9_-]", "_", section_title)[:60]
    short_hash = content_hash[:8]
    return f"crawl_{source_name}_{clean_title}_{short_hash}"


# ─── Ingestion ───────────────────────────────────────────────────

async def _get_crawl_state_collection():
    """Get the crawl_state collection for tracking what's been ingested."""
    from database import _db
    return _db["crawl_state"]


async def _is_already_ingested(content_hash: str) -> bool:
    """Check if a section with this content hash has already been ingested."""
    crawl_state = await _get_crawl_state_collection()
    existing = await crawl_state.find_one({"content_hash": content_hash})
    return existing is not None


async def _mark_as_ingested(content_hash: str, episode_id: str, source_name: str):
    """Record that this content hash has been ingested."""
    crawl_state = await _get_crawl_state_collection()
    await crawl_state.update_one(
        {"content_hash": content_hash},
        {
            "$set": {
                "content_hash": content_hash,
                "episode_id": episode_id,
                "source_name": source_name,
                "ingested_at": datetime.now(timezone.utc),
            }
        },
        upsert=True,
    )


async def _ingest_section(
    source: dict,
    section: dict,
    tags: list[str],
) -> bool:
    """Embed and store a single section as an episode. Returns True if ingested."""

    content_hash = section["hash"]

    # Dedup check
    if await _is_already_ingested(content_hash):
        return False

    episode_id = _make_episode_id(source["name"], section["title"], content_hash)
    collection = get_episodes_collection()

    # Also check if episode_id already exists (belt and suspenders)
    existing = await collection.find_one({"episode_id": episode_id})
    if existing:
        await _mark_as_ingested(content_hash, episode_id, source["name"])
        return False

    # Generate embedding from the section content
    summary = section["content"]
    _set_activity(f"Embedding: {section['title'][:80]}")
    try:
        embedding_vector = get_embedding(summary)
    except Exception as e:
        logger.error(f"Embedding failed for {episode_id}: {e}")
        return False

    now = datetime.now(timezone.utc)
    document = {
        "episode_id": episode_id,
        "instance": source["instance"],
        "project": source["project"],
        "summary": summary,
        "raw_exchange": None,
        "tags": tags + ["crawled", f"source:{source['name']}"],
        "embedding": embedding_vector,
        "timestamp": now,
        "retrieval_count": 0,
        "last_retrieved": None,
    }

    await collection.insert_one(document)
    await _mark_as_ingested(content_hash, episode_id, source["name"])

    # Update live counters so the UI can poll mid-cycle
    _crawler_state["episodes_ingested_last_run"] += 1
    _crawler_state["total_episodes_ingested"] += 1

    logger.info(f"Ingested: {episode_id}")
    return True


# ─── Source Scanning ─────────────────────────────────────────────

async def _scan_source(source: dict) -> int:
    """Scan a single source file and ingest new sections. Returns count ingested."""
    path = source["path"]

    if not os.path.exists(path):
        logger.debug(f"Source not found (skipping): {path}")
        return 0

    try:
        content = Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Failed to read {path}: {e}")
        return 0

    if not content.strip():
        return 0

    # Choose parser based on source type
    if source["name"] == "intercom":
        sections = _parse_intercom_messages(content)
        default_tags = ["intercom"]
    else:
        sections = _parse_markdown_sections(content)
        default_tags = [source["name"]]

    ingested_count = 0
    for section in sections:
        try:
            if await _ingest_section(source, section, default_tags):
                ingested_count += 1
        except Exception as e:
            logger.error(f"Failed to ingest section '{section['title']}' from {source['name']}: {e}")

    return ingested_count


async def _scan_teachings_dir() -> int:
    """Scan the teachings directory for .md files. Each file = one episode."""
    if not os.path.isdir(TEACHINGS_DIR):
        logger.debug(f"Teachings dir not found (skipping): {TEACHINGS_DIR}")
        return 0

    ingested_count = 0
    for root, _dirs, files in os.walk(TEACHINGS_DIR):
        for filename in files:
            if not filename.endswith(".md"):
                continue

            filepath = os.path.join(root, filename)
            try:
                content = Path(filepath).read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                logger.error(f"Failed to read teaching file {filepath}: {e}")
                continue

            if not content.strip() or len(content) < 50:
                continue

            content_hash = _content_hash(content)
            if await _is_already_ingested(content_hash):
                continue

            # Treat the whole file as one episode
            relative_path = os.path.relpath(filepath, TEACHINGS_DIR)
            source = {
                "name": f"teaching_{relative_path}",
                "instance": "office",
                "project": "0_TEACHINGS",
            }
            section = {
                "title": filename.replace(".md", ""),
                "content": content,
                "hash": content_hash,
            }

            try:
                if await _ingest_section(source, section, ["teaching", "learning"]):
                    ingested_count += 1
            except Exception as e:
                logger.error(f"Failed to ingest teaching {filepath}: {e}")

    return ingested_count


async def _scan_documents_dir() -> int:
    """Scan the documents directory for supported document files.

    Supported formats: .docx, .pdf, .odt, .pages, .md, .txt, .rtf
    Each file = one episode.
    """
    if not os.path.isdir(DOCUMENTS_DIR):
        logger.debug(f"Documents dir not found (skipping): {DOCUMENTS_DIR}")
        return 0

    ingested_count = 0
    skipped_empty = 0
    for root, _dirs, files in os.walk(DOCUMENTS_DIR):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            parser = _DOC_PARSERS.get(ext)
            if parser is None:
                continue

            filepath = os.path.join(root, filename)

            # Extract text using the appropriate parser
            content = parser(filepath)
            if not content or len(content) < 30:
                skipped_empty += 1
                continue

            content_hash = _content_hash(content)
            if await _is_already_ingested(content_hash):
                continue

            # Build a summary: filename + content
            doc_title = os.path.splitext(filename)[0]
            summary = f"[Document: {doc_title}]\n\n{content}"

            tags = await _categorize_document(filename, content)
            source = {
                "name": f"doc_{doc_title}",
                "instance": "office",
                "project": "personal_documents",
            }
            section = {
                "title": doc_title,
                "content": summary,
                "hash": content_hash,
            }

            try:
                if await _ingest_section(source, section, tags):
                    ingested_count += 1
            except Exception as e:
                logger.error(f"Failed to ingest document {filepath}: {e}")

    if skipped_empty:
        logger.debug(f"  documents: skipped {skipped_empty} empty/tiny files")
    return ingested_count


async def _scan_projects_dir(machine_root: str, machine_name: str) -> int:
    """Discover projects by docker-compose.yml presence and ingest key files.

    Per project: CLAUDE.md, README.md, docker-compose.yml/yaml, *.sh, *.py
    (all at project root only — no deep recursion into project subdirs).
    """
    root_path = Path(machine_root)
    if not root_path.is_dir():
        logger.debug(f"Machine root not found (skipping): {machine_root}")
        return 0

    # Collect unique project dirs
    project_dirs: set[Path] = set()
    for pattern in ["*/docker-compose.yml", "*/docker-compose.yaml"]:
        for p in root_path.glob(pattern):
            project_dirs.add(p.parent)

    ingested_count = 0
    for project_dir in sorted(project_dirs):
        project_name = project_dir.name

        # Named files to always attempt
        candidates: list[Path] = [
            project_dir / "CLAUDE.md",
            project_dir / "README.md",
            project_dir / "docker-compose.yml",
            project_dir / "docker-compose.yaml",
        ]
        # Code files at project root
        candidates += sorted(project_dir.glob("*.sh"))
        candidates += sorted(project_dir.glob("*.py"))

        for filepath in candidates:
            if not filepath.exists():
                continue

            ext = filepath.suffix.lower()
            source_name = f"{machine_name}_{project_name}_{filepath.name.replace('.', '_')}"
            source = {"name": source_name, "instance": machine_name, "project": project_name}

            if ext == ".md":
                content_text = filepath.read_text(encoding="utf-8", errors="replace")
                sections = _parse_markdown_sections(content_text)
                tags = ["project", machine_name, project_name, "markdown"]
            else:
                sections = _parse_code_file(str(filepath))
                lang = "bash" if ext == ".sh" else ("python" if ext == ".py" else "compose")
                tags = ["project", machine_name, project_name, "code", lang]

            for section in sections:
                try:
                    if await _ingest_section(source, section, tags):
                        ingested_count += 1
                except Exception as e:
                    logger.error(f"Failed to ingest {filepath}: {e}")

    return ingested_count


async def _scan_deep_projects(machine_root: str, machine_name: str) -> int:
    """Scan all 0_* dirs + HUBITAT/NETWORK for code files recursively.

    Discovers projects by directory name (0_* prefix or _EXTRA_PROJECT_NAMES).
    Recurses into subdirs, skipping build/vendor/archive dirs.
    Ingests .ino, .cpp, .h, .groovy, .py, .sh, .js, .ts, .md, etc.
    Files over _MAX_CODE_FILE_BYTES are skipped.
    """
    root_path = Path(machine_root)
    if not root_path.is_dir():
        return 0

    # Find eligible project dirs
    project_dirs: list[Path] = []
    for entry in sorted(root_path.iterdir()):
        if not entry.is_dir():
            continue
        name = entry.name
        if name.startswith("0_") or name in _EXTRA_PROJECT_NAMES:
            project_dirs.append(entry)

    ingested_count = 0
    for project_dir in project_dirs:
        project_name = project_dir.name

        for root, dirs, files in os.walk(project_dir):
            # Prune skip dirs in-place
            dirs[:] = [d for d in dirs if d not in _SKIP_DIRS]

            for filename in files:
                filepath = os.path.join(root, filename)
                ext = os.path.splitext(filename)[1].lower()

                if ext == ".md":
                    pass  # markdown is always welcome
                elif ext not in _DEEP_CODE_EXTENSIONS:
                    continue

                # Skip oversized files
                try:
                    if os.path.getsize(filepath) > _MAX_CODE_FILE_BYTES:
                        continue
                except OSError:
                    continue

                rel_path = os.path.relpath(filepath, str(project_dir))
                source_name = f"{machine_name}_{project_name}_{re.sub(r'[^a-zA-Z0-9_-]', '_', rel_path)}"
                source = {"name": source_name, "instance": machine_name, "project": project_name}

                if ext == ".md":
                    try:
                        content = Path(filepath).read_text(encoding="utf-8", errors="replace")
                    except Exception:
                        continue
                    sections = _parse_markdown_sections(content)
                    tags = ["project", machine_name, project_name, "markdown"]
                else:
                    sections = _parse_code_file(filepath)
                    # Determine language tag
                    lang_map = {
                        ".ino": "arduino", ".cpp": "cpp", ".c": "c", ".h": "c-header",
                        ".groovy": "groovy", ".py": "python", ".sh": "bash",
                        ".js": "javascript", ".ts": "typescript",
                        ".jsx": "react", ".tsx": "react",
                        ".yml": "yaml", ".yaml": "yaml",
                        ".java": "java", ".go": "go", ".rs": "rust",
                    }
                    lang = lang_map.get(ext, "code")
                    tags = ["project", machine_name, project_name, "code", lang]

                for section in sections:
                    try:
                        if await _ingest_section(source, section, tags):
                            ingested_count += 1
                    except Exception as e:
                        logger.error(f"Failed to ingest {filepath}: {e}")

    if ingested_count > 0:
        logger.info(f"  {machine_name} deep projects: {ingested_count} new episodes")
    return ingested_count


async def _scan_scripts_dir(machine_root: str, machine_name: str) -> int:
    """Scan 0_SCRIPTS/ for .sh files (bash style authority).

    Skips dirs containing DEPRECATED, TRASH, ARCHIVE, CONFLICTS.
    Each .sh file = one episode (whole file, no section splitting).
    """
    scripts_dir = Path(machine_root) / "0_SCRIPTS"
    if not scripts_dir.is_dir():
        logger.debug(f"Scripts dir not found (skipping): {scripts_dir}")
        return 0

    ingested_count = 0
    for root, dirs, files in os.walk(scripts_dir):
        # Prune skip dirs in-place
        dirs[:] = [
            d for d in dirs
            if not any(skip in d.upper() for skip in _SCRIPTS_SKIP)
        ]
        for filename in files:
            if not filename.endswith(".sh"):
                continue
            filepath = os.path.join(root, filename)
            sections = _parse_code_file(filepath)
            if not sections:
                continue
            rel_path = os.path.relpath(filepath, str(scripts_dir))
            source_name = f"{machine_name}_scripts_{re.sub(r'[^a-zA-Z0-9_-]', '_', rel_path)}"
            source = {"name": source_name, "instance": machine_name, "project": "0_SCRIPTS"}
            tags = ["script", "bash", "code-style", machine_name]
            for section in sections:
                try:
                    if await _ingest_section(source, section, tags):
                        ingested_count += 1
                except Exception as e:
                    logger.error(f"Failed to ingest script {filepath}: {e}")

    return ingested_count


# ─── Main Crawl Loop ────────────────────────────────────────────

async def run_crawl_cycle():
    """Execute one full crawl cycle across all sources."""
    _crawler_state["running"] = True
    _crawler_state["episodes_ingested_last_run"] = 0
    start_time = datetime.now(timezone.utc)
    total_ingested = 0
    errors = []

    logger.info("Crawl cycle starting...")
    _set_activity("Starting crawl cycle...")

    # Scan each defined source
    for source in SOURCES:
        _set_activity(f"Scanning source: {source['name']}")
        try:
            count = await _scan_source(source)
            total_ingested += count
            if count > 0:
                logger.info(f"  {source['name']}: {count} new episodes")
        except Exception as e:
            error_msg = f"{source['name']}: {e}"
            logger.error(f"  {error_msg}")
            errors.append(error_msg)

    # Scan each machine's projects and scripts
    for machine_name, machine_root in MACHINE_ROOTS.items():
        _set_activity(f"Scanning {machine_name} projects...")
        try:
            count = await _scan_projects_dir(machine_root, machine_name)
            total_ingested += count
            if count > 0:
                logger.info(f"  {machine_name} projects: {count} new episodes")
        except Exception as e:
            error_msg = f"{machine_name}_projects: {e}"
            logger.error(f"  {error_msg}")
            errors.append(error_msg)

        _set_activity(f"Scanning {machine_name} scripts...")
        try:
            count = await _scan_scripts_dir(machine_root, machine_name)
            total_ingested += count
            if count > 0:
                logger.info(f"  {machine_name} scripts: {count} new episodes")
        except Exception as e:
            error_msg = f"{machine_name}_scripts: {e}"
            logger.error(f"  {error_msg}")
            errors.append(error_msg)

        _set_activity(f"Deep scanning {machine_name} (0_* dirs)...")
        try:
            count = await _scan_deep_projects(machine_root, machine_name)
            total_ingested += count
        except Exception as e:
            error_msg = f"{machine_name}_deep_projects: {e}"
            logger.error(f"  {error_msg}")
            errors.append(error_msg)

    # Scan teachings directory
    _set_activity("Scanning teachings...")
    try:
        teaching_count = await _scan_teachings_dir()
        total_ingested += teaching_count
        if teaching_count > 0:
            logger.info(f"  teachings: {teaching_count} new episodes")
    except Exception as e:
        error_msg = f"teachings: {e}"
        logger.error(f"  {error_msg}")
        errors.append(error_msg)

    # Scan documents directory
    _set_activity("Scanning documents (.docx, .pdf, .odt, .pages, .md, .txt)...")
    try:
        doc_count = await _scan_documents_dir()
        total_ingested += doc_count
        if doc_count > 0:
            logger.info(f"  documents: {doc_count} new episodes")
    except Exception as e:
        error_msg = f"documents: {e}"
        logger.error(f"  {error_msg}")
        errors.append(error_msg)

    duration = (datetime.now(timezone.utc) - start_time).total_seconds()

    _crawler_state["last_run"] = start_time.isoformat()
    _crawler_state["last_run_duration_seconds"] = round(duration, 2)
    # episodes_ingested_last_run and total_episodes_ingested already
    # incremented live inside _ingest_section()
    _crawler_state["errors"] = errors
    _crawler_state["running"] = False
    _set_activity("")

    # Persist stats to MongoDB (survives container restarts)
    coll = get_settings_collection()
    await coll.update_one(
        {"_id": "crawler_stats"},
        {"$set": {
            "total_episodes_ingested": _crawler_state["total_episodes_ingested"],
            "episodes_ingested_last_run": _crawler_state["episodes_ingested_last_run"],
            "last_run": _crawler_state["last_run"],
            "last_run_duration_seconds": _crawler_state["last_run_duration_seconds"],
        }},
        upsert=True,
    )

    logger.info(
        f"Crawl cycle complete: {total_ingested} episodes ingested "
        f"in {duration:.1f}s ({len(errors)} errors)"
    )

    return total_ingested


# start_crawler / stop_crawler removed — crawler is now driven by the
# unified scheduler (scheduler.py) like JSONL and training.
# Manual runs still work via POST /api/crawler/run.
